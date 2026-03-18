"""Build homework documents (LaTeX/PDF + Word) for both assignments."""

import subprocess
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from config import SCREENSHOTS_DIR, OUTPUT_DIR, check_banned_words

# ─── Paragraph text ───────────────────────────────────────────────

AI_PARAGRAPH = (
    "The Ministry of Education runs China's AI education system from top to bottom. "
    "It sets curricula at every school level and has pushed every major AI education "
    "directive since 2018, starting with the AI Innovation Action Plan that created "
    '"AI+X" programs at universities and expanded doctoral enrollment. In late 2024, '
    "the MOE told all K-12 schools to teach AI across every grade and picked 184 schools "
    "as pilot bases. By September 2025, every primary and secondary school in China had "
    "to offer at least 8 hours of AI instruction per year."
)

LANGUAGE_PARAGRAPH = (
    "The Ministry of Education controls language education at every level in China, "
    "from primary school English classes to the gaokao exam. English stays one of three "
    "required gaokao subjects, worth 150 out of 750 points, and students start learning "
    "it in third grade. The MOE also runs the State Language Commission and oversees the "
    "Center for Language Education and Cooperation. Its 2021 \"Double Reduction\" ban on "
    "for-profit K-9 tutoring reshaped how millions of students access language instruction, "
    "while recent gaokao reforms now let students sit the English exam twice a year."
)

# ─── Document metadata ────────────────────────────────────────────

AI_DOC = {
    "filename_base": "ai_education_homework",
    "title": "China's Government Institutions and AI Education Policy",
    "student_name": "Your Name",  # User's doc
    "research_topic": "How does the Chinese government regulate and promote artificial intelligence education?",
    "primary_institution": "Ministry of Education (教育部)",
    "primary_website": "https://www.moe.gov.cn/",
    "primary_screenshot": "moe_gov_cn.png",
    "secondary_institutions": [
        "Ministry of Science and Technology (科学技术部) — https://www.most.gov.cn/",
        "Cyberspace Administration of China (国家互联网信息办公室) — https://www.cac.gov.cn/",
    ],
    "paragraph": AI_PARAGRAPH,
}

LANGUAGE_DOC = {
    "filename_base": "language_education_homework",
    "title": "China's Government Institutions and Language Education Policy",
    "student_name": "Gaby",  # Gaby's doc
    "research_topic": "How does the Chinese government manage and shape language education policy?",
    "primary_institution": "Ministry of Education (教育部)",
    "primary_website": "https://www.moe.gov.cn/",
    "primary_screenshot": "moe_gov_cn.png",
    "secondary_institutions": [
        "State Language Commission (国家语言文字工作委员会) — http://www.moe.gov.cn/jyb_sy/China_Language/",
        "Center for Language Education and Cooperation (中外语言交流合作中心) — https://www.chinese.cn/",
    ],
    "paragraph": LANGUAGE_PARAGRAPH,
}

# ─── LaTeX template ───────────────────────────────────────────────

LATEX_TEMPLATE = r"""\documentclass[12pt,a4paper]{{article}}
\usepackage{{xeCJK}}
\setCJKmainfont{{Droid Sans Fallback}}
\usepackage{{geometry}}
\geometry{{margin=1in}}
\usepackage{{graphicx}}
\usepackage{{hyperref}}
\usepackage{{enumitem}}
\usepackage{{setspace}}
\onehalfspacing

\title{{{title}}}
\author{{{student_name}}}
\date{{\today}}

\begin{{document}}
\maketitle

\section*{{Research Topic}}
{research_topic}

\section*{{Primary Institution}}
\textbf{{{primary_institution}}}\\
Website: \url{{{primary_website}}}

\begin{{figure}}[h]
\centering
\includegraphics[width=0.85\textwidth]{{{screenshot_path}}}
\caption{{Screenshot of {primary_institution} website}}
\end{{figure}}

\section*{{Other Relevant Institutions}}
\begin{{itemize}}
{secondary_items}
\end{{itemize}}

\section*{{Summary Paragraph}}
{paragraph}

\end{{document}}
"""


def validate_paragraphs():
    """Check both paragraphs for banned words and word count."""
    errors = []
    for label, text in [("AI", AI_PARAGRAPH), ("Language", LANGUAGE_PARAGRAPH)]:
        banned = check_banned_words(text)
        if banned:
            errors.append(f"{label} paragraph contains banned words: {banned}")

        word_count = len(text.split())
        if word_count < 80 or word_count > 100:
            errors.append(f"{label} paragraph word count: {word_count} (need 80-100)")
        else:
            print(f"  {label} paragraph: {word_count} words — OK")

    if errors:
        for e in errors:
            print(f"  ERROR: {e}")
        sys.exit(1)


def build_latex(doc: dict):
    """Generate LaTeX source and compile to PDF."""
    base = doc["filename_base"]
    screenshot_path = SCREENSHOTS_DIR / doc["primary_screenshot"]

    # Build secondary institution items
    items = "\n".join(f"  \\item {inst}" for inst in doc["secondary_institutions"])

    # LaTeX needs forward slashes and no backslash issues
    screenshot_str = str(screenshot_path).replace("\\", "/")

    latex_src = LATEX_TEMPLATE.format(
        title=doc["title"],
        student_name=doc["student_name"],
        research_topic=doc["research_topic"],
        primary_institution=doc["primary_institution"],
        primary_website=doc["primary_website"],
        screenshot_path=screenshot_str,
        secondary_items=items,
        paragraph=doc["paragraph"],
    )

    tex_path = OUTPUT_DIR / f"{base}.tex"
    tex_path.write_text(latex_src, encoding="utf-8")
    print(f"  Written: {tex_path.name}")

    # Compile with xelatex
    print(f"  Compiling {base}.pdf...")
    result = subprocess.run(
        ["xelatex", "-interaction=nonstopmode", f"-output-directory={OUTPUT_DIR}", str(tex_path)],
        capture_output=True,
        text=True,
        timeout=120,
    )

    pdf_path = OUTPUT_DIR / f"{base}.pdf"
    if pdf_path.exists():
        size_kb = pdf_path.stat().st_size / 1024
        print(f"  Compiled: {pdf_path.name} ({size_kb:.0f} KB)")
    else:
        print(f"  WARNING: PDF compilation failed")
        # Show last 20 lines of log for debugging
        log_path = OUTPUT_DIR / f"{base}.log"
        if log_path.exists():
            lines = log_path.read_text(errors="replace").split("\n")
            error_lines = [l for l in lines if "!" in l or "Error" in l]
            for l in error_lines[:10]:
                print(f"    {l}")


def build_word(doc: dict):
    """Generate Word document."""
    base = doc["filename_base"]
    screenshot_path = SCREENSHOTS_DIR / doc["primary_screenshot"]

    document = Document()

    # Title
    title_para = document.add_heading(doc["title"], level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Student name
    name_para = document.add_paragraph(doc["student_name"])
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in name_para.runs:
        run.font.size = Pt(12)

    document.add_paragraph()  # Spacer

    # Research Topic
    document.add_heading("Research Topic", level=1)
    document.add_paragraph(doc["research_topic"])

    # Primary Institution
    document.add_heading("Primary Institution", level=1)
    p = document.add_paragraph()
    run = p.add_run(doc["primary_institution"])
    run.bold = True
    document.add_paragraph(f"Website: {doc['primary_website']}")

    # Screenshot
    if screenshot_path.exists():
        document.add_picture(str(screenshot_path), width=Inches(5.5))
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        document.add_paragraph("[Screenshot not available]")

    # Secondary institutions
    document.add_heading("Other Relevant Institutions", level=1)
    for inst in doc["secondary_institutions"]:
        document.add_paragraph(inst, style="List Bullet")

    # Summary paragraph
    document.add_heading("Summary Paragraph", level=1)
    para = document.add_paragraph(doc["paragraph"])
    for run in para.runs:
        run.font.size = Pt(11)
        run.font.name = "Calibri"

    docx_path = OUTPUT_DIR / f"{base}.docx"
    document.save(str(docx_path))
    size_kb = docx_path.stat().st_size / 1024
    print(f"  Written: {docx_path.name} ({size_kb:.0f} KB)")


def main():
    print("=" * 60)
    print("Building homework documents")
    print("=" * 60)

    # Validate paragraphs first
    print("\nValidating paragraphs...")
    validate_paragraphs()

    for doc in [AI_DOC, LANGUAGE_DOC]:
        print(f"\n--- {doc['title']} ---")

        print("\n  Building LaTeX/PDF...")
        build_latex(doc)

        print("\n  Building Word document...")
        build_word(doc)

    # Clean up LaTeX auxiliary files
    for ext in ["aux", "log", "out"]:
        for f in OUTPUT_DIR.glob(f"*.{ext}"):
            f.unlink()

    print("\n" + "=" * 60)
    print("Build complete!")
    print("=" * 60)
    print(f"\nOutput files in: {OUTPUT_DIR}")
    for f in sorted(OUTPUT_DIR.iterdir()):
        print(f"  {f.name} ({f.stat().st_size / 1024:.0f} KB)")


if __name__ == "__main__":
    main()
